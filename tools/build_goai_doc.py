from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\shiwenbo\.mavis\agents\mavis\workspace\llm-reputation-paper\llm-reputation")
OUT = ROOT / "LLM驱动开放式策略演化平台_初赛问题定义文档_更新版.docx"

NAVY = "183153"
BLUE = "2563A6"
TEAL = "177E89"
LIGHT = "EAF2F8"
PALE = "F4F7FA"
GRAY = "536273"
WHITE = "FFFFFF"
RED = "A23B3B"


def font(run, size=9, bold=False, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=110, bottom=80, end=110):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, color=None, size=8.4, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    font(r, size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def add_title(doc, page_no, section_title, kicker):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(15.7)
    table.columns[1].width = Cm(2.0)
    table.cell(0, 0).width = Cm(15.7)
    table.cell(0, 1).width = Cm(2.0)
    set_cell_text(table.cell(0, 0), "GOAI · AI for Research｜开放探索赛初赛问题定义", True, WHITE, 9.5)
    set_cell_text(table.cell(0, 1), f"PAGE {page_no}/4", True, WHITE, 9, WD_ALIGN_PARAGRAPH.CENTER)
    shade(table.cell(0, 0), NAVY)
    shade(table.cell(0, 1), BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(section_title)
    font(r, 18, True, NAVY)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(kicker)
    font(r, 8.5, False, GRAY)


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    font(r, 10.8, True, BLUE)
    return p


def add_body(doc, text, size=8.8, after=2, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        font(r, size, True, NAVY)
        r = p.add_run(text[len(bold_lead):])
        font(r, size)
    else:
        r = p.add_run(text)
        font(r, size)
    return p


def add_bullet(doc, text, size=8.6):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.55)
    p.paragraph_format.first_line_indent = Cm(-0.28)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    font(r, size)


def add_callout(doc, label, text, fill=LIGHT, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table.cell(0, 0)
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label + "  ")
    font(r, 8.6, True, color)
    r = p.add_run(text)
    font(r, 8.6)
    margins(c, 100, 150, 100, 150)
    return table


def add_figure(doc, path, caption, width_cm, height_cm=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    if height_cm:
        p.add_run().add_picture(str(path), width=Cm(width_cm), height=Cm(height_cm))
    else:
        p.add_run().add_picture(str(path), width=Cm(width_cm))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(caption)
    font(r, 7.5, False, GRAY)


def add_page_meta(doc, evidence, visual, confirm, budget):
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    values = [
        ("证据编号", evidence), ("建议视觉", visual),
        ("人工确认", confirm), ("内容预算", budget),
    ]
    for idx, (label, value) in enumerate(values):
        c = table.cell(idx // 2, idx % 2)
        c.text = ""
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label + "｜")
        font(r, 7.2, True, TEAL)
        r = p.add_run(value)
        font(r, 7.2, False, GRAY)
        shade(c, PALE)
        margins(c, 65, 100, 65, 100)


def add_source_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    font(r, 6.8, False, GRAY)


doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(1.0)
sec.bottom_margin = Cm(1.0)
sec.left_margin = Cm(1.35)
sec.right_margin = Cm(1.35)
sec.header_distance = Cm(0.35)
sec.footer_distance = Cm(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Microsoft YaHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
normal.font.size = Pt(8.8)
normal.paragraph_format.space_after = Pt(2)
for name in ("Heading 1", "Heading 2", "Heading 3"):
    st = styles[name]
    st.font.name = "Microsoft YaHei"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

footer = sec.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LLM 驱动开放式策略演化平台｜问题定义更新版｜2026-08")
font(r, 7, False, GRAY)

# PAGE 1
add_title(doc, 1, "一、问题与证据", "从“通用策略竞技平台”收敛为：私有声誉条件下，LLM 生成的可执行评价算法能否通过演化维持合作？")
add_callout(doc, "四页共用信息", "项目标题：基于 LLM 驱动代码演化的私有声誉算法研究｜边界：15 个智能体的捐赠博弈、私有声誉库、受控可观测率与固定收益结构｜成熟度：核心环境、主实验、对照和入侵测试均已实现；论文与复现实验仍在迭代。")
add_heading(doc, "1.1 真实问题或需求")
add_body(doc, "间接互惠依赖声誉：个体会根据他人过去对第三方的行为决定是否合作。经典模型往往预设公开、同步的声誉或少数人工规范；但现实中的在线社区和多智能体系统通常只有碎片化观察，每个主体独立形成评价。已有研究表明，私有且含噪的信息可能造成评价分裂，而更细粒度的定量声誉又可能恢复合作［E1–E2］。", 8.7)
add_heading(doc, "1.2 为什么尚未被充分结构化")
add_bullet(doc, "理论工作通常先指定 Image Scoring、Standing 或 leading-eight 等规则，再分析其稳定性；“允许策略从原始观察中自由形成时会出现什么算法”仍缺少统一基准。")
add_bullet(doc, "现有 LLM 社会博弈研究多让模型逐轮行动，或在双人重复博弈中生成完整策略；私有声誉、N 人互动、策略代码演化与事后机制分析尚未被同一环境整合［E3］。")
add_bullet(doc, "仅看平均合作率会混淆无条件合作、可抵抗背叛的合作与偶然路径；必须同时定义对照、跨种子稳定性和入侵检验。")
add_heading(doc, "1.3 研究价值与合适切片")
add_body(doc, "本项目不追求“找到永远最强策略”，而把科学问题切为：在固定收益成本比、私有声誉更新和不同观察概率 p 下，LLM 作为语义变异算子能否提出可执行的评价/行动规则，并经演化选择形成可解释、可复现、可受反例挑战的合作机制。AI 的角色是扩展候选算法空间；博弈模拟、适应度和对照实验负责筛选。")
add_figure(doc, ROOT / "results/figures/fig2_observability_scan.png", "项目现有证据：合作结果随可观测率和随机种子变化；图用于说明“问题可测”，不作为预设答案。", 13.2, 4.1)
add_page_meta(doc, "E1、E2、E3、E4", "可观测率—合作率图", "确认 E1–E3 原始页面及项目边界表述", "正文 650–800 字；1 图")
add_source_note(doc, "E1 Schmid et al., Nature Communications 14, 2086 (2023), doi:10.1038/s41467-023-37817-x；E2 Ohtsuki & Iwasa, JTB 239 (2006), doi:10.1016/j.jtbi.2005.08.008；E3 Willis et al., arXiv:2501.16173；E4 项目 README 与 paper_zh/paper.tex（项目内核验）。")

doc.add_page_break()

# PAGE 2
add_title(doc, 2, "二、环境接口", "Agent 改变的是可执行策略代码；反馈来自同一规则下的交互收益、合作轨迹、选择事件和跨策略检验。")
add_heading(doc, "2.1 固定规则")
add_body(doc, "环境为 N=15 的捐赠博弈。每次交互中，捐赠者合作需付出成本 c=1，接收者获得收益 b=2；个体适应度由收到收益减去付出成本累计。每个智能体维护独立的 `reputations: dict[int,float]`，无中央声誉。可观测率 p 控制每个第三方观察被看到的概率；博弈规则、收益、种群规模、每代交互预算、选择方案和沙箱约束在同一实验内固定。")
add_figure(doc, ROOT / "results/figures/framework/framework.png", "已实现的三层闭环：博弈环境 → 私有声誉与策略执行 → LLM 生成/变异与演化选择。", 14.6, 6.0)
add_heading(doc, "2.2 观察 / 行动 / 反馈")
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [2.4, 5.1, 5.1, 5.1]
for i, w in enumerate(widths):
    table.columns[i].width = Cm(w)
headers = ["接口", "具体输入", "Agent 可改变对象", "即时/延迟反馈"]
for i, h in enumerate(headers):
    set_cell_text(table.cell(0, i), h, True, WHITE, 7.9, WD_ALIGN_PARAGRAPH.CENTER)
    shade(table.cell(0, i), NAVY)
rows = [
    ("观察", "目标当前声誉、目标动作、自身声誉；LLM 变异时另见父代源码与适应度", "不改变固定规则；读取规定字段", "观察是否成功写入私有声誉库"),
    ("评价", "`evaluate(target_rep, target_action, my_rep)`", "声誉更新公式、增量、阈值与条件分支", "返回值、越界裁剪、执行异常"),
    ("决策", "`decide(my_rep, opponent_rep)`", "合作阈值及对自身/对手声誉的使用方式", "合作/背叛、单次收益、对手响应"),
    ("演化", "个体适应度、父代源码、验证错误", "生成新策略代码；选择保留/替换的谱系", "代际合作率、适应度、存活、复制和变异事件"),
]
for row in rows:
    cells = table.add_row().cells
    for i, txt in enumerate(row):
        set_cell_text(cells[i], txt, i == 0, NAVY if i == 0 else None, 7.5)
        shade(cells[i], LIGHT if i == 0 else WHITE)

add_heading(doc, "2.3 记录与预算")
add_body(doc, "每次运行保存配置（模型、提示版本、随机种子、p、N、代数、交互数、选择参数）、每代合作率与个体适应度、完整策略源码、私有声誉状态、亲缘/替换事件、验证失败与 fallback 原因。正式比较必须使用相同交互和 LLM 调用预算；生成代码先静态校验，再在受限沙箱执行。当前主实验采用 4 个 p 条件×10 个种子；生产长程实验另使用 100 代、每代 1,000 次交互进行稳定性观察［E4–E5］。")
add_page_meta(doc, "E4、E5", "三层框架图＋接口表", "核对最终提交所引用的配置版本", "正文 550–700 字；1 图 1 表")
add_source_note(doc, "E5 项目代码 experiments/game、experiments/v2_quantitative、experiments/evolution、experiments/sandbox 及结果 JSON（项目内核验）。")

doc.add_page_break()

# PAGE 3
add_title(doc, 3, "三、发现信号与参照", "发现不是“分数更高”本身，而是跨种子出现、可由源码解释、并能在公平参照与反事实测试中保留的机制。")
add_heading(doc, "3.1 什么算发现")
add_bullet(doc, "正向机制：演化策略在未提示经典规范名称的条件下形成可解释的声誉更新与条件合作结构，并在多随机种子或独立复现中重复出现。")
add_bullet(doc, "异常/反例：合作率对可观测率非单调、种子间双峰、复杂代码退化成平凡行为，或经典规范在特定私有信息条件下失效。")
add_bullet(doc, "稳定负结果：LLM 变异不能超过随机或手工基线、策略无法抵抗入侵、结果依赖动作标签/提示措辞，均用于收缩“LLM 能发现机制”的主张。")
add_bullet(doc, "问题修正：若高合作仅来自 ALLC、fallback 或评估漏洞，则把目标从“合作涌现”改为“可抵抗利用且可解释的声誉机制”。")
add_heading(doc, "3.2 平凡解 / 随机 / 无干预 / 非平凡基线")
table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
headers = ["参照", "排除的伪解释", "比较指标", "超过参照的条件"]
for i, h in enumerate(headers):
    set_cell_text(table.cell(0, i), h, True, WHITE, 7.9, WD_ALIGN_PARAGRAPH.CENTER)
    shade(table.cell(0, i), NAVY)
baseline_rows = [
    ("ALLC / ALLD", "高合作或高个体收益只是平凡规则", "合作率、适应度、被入侵率", "不只群体内高合作，还能抵抗剥削"),
    ("随机策略/随机变异", "搜索收益来自运气或演化本身", "跨种子分布、存活率、最终代码族", "LLM 语义变异稳定优于随机变异"),
    ("静态无演化", "初始模型偏好被误认成演化发现", "初代与末代差异、选择谱系", "选择后出现可重复的结构/行为变化"),
    ("经典规范", "只重新命名已知规则", "对局矩阵、入侵方向、机制差异", "相对 Image Scoring / leading-eight 有可解释增益或边界发现"),
]
for row in baseline_rows:
    cells = table.add_row().cells
    for i, txt in enumerate(row):
        set_cell_text(cells[i], txt, i == 0, NAVY if i == 0 else None, 7.4)
        shade(cells[i], LIGHT if i == 0 else WHITE)

add_heading(doc, "3.3 最低成功与失败标准")
add_callout(doc, "最低成功", "至少在 3 个随机种子中出现同类可解释策略结构；相对随机变异和静态无演化对照有方向一致的差异；原始源码、配置与轨迹完整；至少通过一种未参与选择的外部检验（隐藏对手、参数外推或入侵测试）。")
add_callout(doc, "明确失败", "若优势只存在于单一种子/同代对手，去除 fallback 后消失，换动作标签即翻转，或在外部检验中不优于 ALLC/随机/经典规范，则不宣称新机制；保留为提示偏差、评估漏洞或适用边界的负结果。", "FBECEC", RED)
add_figure(doc, ROOT / "results/figures/fig4_selection_comparison.png", "现有对照证据：选择与 LLM 语义变异的贡献需要拆开比较；正式结论以多种子分布和源码审计为准。", 12.6, 3.8)
add_page_meta(doc, "E5、E6", "对照结果图或双向入侵仪表盘", "确认最低阈值是否沿用至复赛", "正文 600–750 字；1 图 1 表")
add_source_note(doc, "E6 入侵实验 summary.json：4 个经典规范×2 个方向×14 个初始数量×3 个种子，共 336 次固定策略运行（项目内核验）。")

doc.add_page_break()

# PAGE 4
add_title(doc, 4, "四、最小验证计划", "首版闭环已经跑通；下一步不是重复“再跑一次”，而是用可审计的复现实验确认机制、边界和失败原因。")
add_heading(doc, "4.1 一次可复现试跑怎么做")
steps = [
    ("01｜锁定输入", "固定代码提交、配置、提示版本、模型端点、seed、p、N、代数和每代交互预算；先执行 dry-run 与沙箱自检。"),
    ("02｜生成与校验", "LLM 生成初始 `evaluate/decide`；记录原始响应、编译/接口校验、超时与 fallback，不允许静默替换。"),
    ("03｜环境评估", "在同一捐赠博弈内计算个体适应度、合作率、私有声誉分歧和行为矩阵；保存逐代完整 JSON。"),
    ("04｜选择与变异", "按固定选择机制保留/复制策略，LLM 对父代源码作语义变异；写入 parent_id、mutation_event 和代码哈希。"),
    ("05｜外部检验", "重复多种子；与 ALLC/ALLD、随机变异、静态无演化和经典规范比较；对最终候选运行双向入侵测试。"),
    ("06｜机制归纳", "对最终源码去重、聚类和人工复核；把行为标签与原始代码绑定，主动寻找反例，不以单条漂亮轨迹下结论。"),
]
for label, text in steps:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_cell_text(table.cell(0, 0), label, True, WHITE, 7.8, WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(0, 1), text, False, None, 7.7)
    shade(table.cell(0, 0), TEAL)
    shade(table.cell(0, 1), PALE)

add_heading(doc, "4.2 主要风险与失败路径")
risks = [
    ("生成失败 / API 限流", "显式记录失败率与回退来源；若某条件 fallback 占比异常，则该批结果无效并重跑。"),
    ("提示泄漏与词汇偏差", "使用无理论名称提示；做 cooperate/defect、donate/not_donate、A/B 标签消融。"),
    ("单种子叙事与评估过拟合", "预先固定种子数、报告分布和轨迹；用隐藏对手、参数外推及双向入侵检验。"),
    ("沙箱漏洞 / 行为—源码错配", "限制导入、长度与执行时间；保存代码哈希，抽样回放策略输出并审计无效分支。"),
]
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(("风险", "定位、留证与调整")):
    set_cell_text(table.cell(0, i), h, True, WHITE, 7.9, WD_ALIGN_PARAGRAPH.CENTER)
    shade(table.cell(0, i), NAVY)
for a, b in risks:
    cells = table.add_row().cells
    set_cell_text(cells[0], a, True, NAVY, 7.6)
    set_cell_text(cells[1], b, False, None, 7.6)
    shade(cells[0], LIGHT)

add_heading(doc, "4.3 复现与开源计划")
add_body(doc, "仓库已按环境、演化、Agent、沙箱、分析和结果分层；使用 Python 3.12、`uv.lock` 锁定依赖，并提供 dry-run、单/多 seed 运行、图表生成和入侵实验入口。提交版计划公开：去密钥配置、提示模板、策略接口、核心实验脚本、随机种子、聚合结果与代表性原始 JSON；API 密钥和可能受服务条款限制的原始响应不公开。复现者应先运行 smoke test，再执行主实验与图表校验。")
add_page_meta(doc, "E4、E5、E6", "六步闭环＋风险表", "移除密钥/个人信息；确定公开许可与仓库链接", "正文 650–800 字；2 表")
add_source_note(doc, "证据状态说明：E1–E3 已由本次编辑打开原始出版/预印本页面，提交前仍建议作者人工确认；E4–E6 已对照当前仓库实现与结果文件。本文档中的实验结果用于说明当前成熟度，不倒灌为事前成功标准。")

doc.save(OUT)
print(OUT)
